from typing import List, Dict
from docxtpl import DocxTemplate
from docx import Document
from io import BytesIO
from Models.FinalContract import FinalContract
from Services.DataSource import DataSource
from Utilities.CreationResult import CreationResult

#---------------------------DATASERVICE CLASS---------------------------
#Clase principal para la automatizacion
class DataService:
    #----------------------attrs-----------------------
    def __init__(self, data_src: DataSource):
        self._template_path: str = "DataSources\\template.docx"
        self._vars: List[str] = [
            "ctoNumber", "ctoClass", "ctoObject", "ctoInitialValue", "ctoAddedValue", "ctoFinalValue", 
            "employeeName", "employeeDni", "employeeOverseerGenre", "employeeOverseerName"
        ]
        self._data_src: DataSource = data_src
        
    #----------------------methods-----------------------
    #----------------------load meths-----------------------
    def load_contracts(self) -> List[FinalContract] | None:
        try:
            return self._data_src.get_final_cts()
        except Exception as e:
            print(f"There has been an error loading contracts data DataService-load_contracts: {e}")
            return None

    #----------------------get meths-----------------------
    def get_vars_list(self) -> List[str]:
        return self._vars
    
    def get_docx_as_html_txt(self) -> str | None:
        try:
            doc = Document(self._template_path)
            content: str = ""
            
            for parag in doc.paragraphs:
                content += f"<p>{parag.text}<p>"
            
            return content
        except Exception as e:
            print(f"There has been an error loading file to show data DataService-get_docx_to_show: {e}")
            return None
    
    def get_docx_bytes(self) -> BytesIO | None:
        try:
            doc = Document(self._template_path)
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer
        except Exception as e:
            print(f"There has been an error geting binary docx file DataService-get_docx_binary: {e}")
            return None
    
    #----------------------creation meths-----------------------
    def create_docx_file(self, contract: FinalContract, path: str) -> CreationResult:
        try:
            doc = DocxTemplate(self._template_path)

            context: Dict[str, str] = {
                var: str(contract.get_attr_by_var_name(var)) for var in self._vars
            }

            doc.render(context)

            contract_name: str = f"OTROSI {contract.number}-2025 - {contract.emp_role} {contract.emp_name}.docx"

            doc.save(f"{path}\\{contract_name}")
            return CreationResult.ok(f"{contract_name} fue creado exitosamente")

        except Exception as e:
            return CreationResult.fail("Error creando archivo .docx en DataService-create_docx_file", str(e))
        